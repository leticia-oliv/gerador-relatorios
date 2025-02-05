import pandas as pd
from docx import Document
from tkinter import messagebox

def gerar_relatorio(arquivo_excel, pasta_destino):
    if not arquivo_excel or not pasta_destino:
        messagebox.showerror("Erro", "Selecione um arquivo Excel e uma pasta de destino!")
        return

    df = pd.read_excel(arquivo_excel)
    for nome, grupo in df.groupby('Nome'):
        doc = Document()
        doc.add_paragraph("Relatório de Gastos", style="Title")
        doc.add_paragraph(f"Nome: {nome}", style="Heading 1")
        doc.add_paragraph("Este relatório contém os registros de gastos da planilha.", style="Normal")

        tabela = doc.add_table(rows=1, cols=2)
        tabela.style = "Table Grid"
        hdr_cells = tabela.rows[0].cells
        hdr_cells[0].text = "Débitos"
        hdr_cells[1].text = "Valor (R$)"

        for _, row in grupo.iterrows():
            row_cells = tabela.add_row().cells
            row_cells[0].text = str(row["Débitos"])
            row_cells[1].text = f'R$ {row["Valor"]:.2f}'

        total_gastos = grupo['Valor'].sum()
        doc.add_paragraph(f'\nTOTAL DE GASTOS: R$ {total_gastos:.2f}', style="Heading 2")

        nome_limpo = nome.replace(" ", "_").replace("ç", "c").replace("ã", "a")
        arquivo_word = f"{pasta_destino}/relatorio_gastos_{nome_limpo}.docx"
        doc.save(arquivo_word)

    messagebox.showinfo("Sucesso", "Relatórios gerados com sucesso!")